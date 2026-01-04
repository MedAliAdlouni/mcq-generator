import sys
from pathlib import Path
from unittest.mock import patch
import pytest

from app import create_app
from app.db import Base, LocalSession
from app.models import User, Document, Question, QuestionType

ROOT_DIR = Path(__file__).resolve().parent.parent
sys.path.append(str(ROOT_DIR))

# ------------------------------------------------
# Test Configuration
# ------------------------------------------------

@pytest.fixture(scope="session")
def test_db_url():
    """Generate a temporary SQLite database URL for testing"""
    # Use in-memory SQLite for fast tests
    return "sqlite:///:memory:"


@pytest.fixture
def app(test_db_url, monkeypatch):
    """Create and configure a Flask application instance for testing"""
    # Set test environment variables
    monkeypatch.setenv("SECRET_KEY", "test-secret-key-for-testing-only")
    monkeypatch.setenv("DATABASE_URL", test_db_url)
    monkeypatch.setenv("FLASK_ENV", "testing")
    monkeypatch.setenv("GEMINI_API_KEY", "test-api-key")  # Mock API key
    
    app = create_app()
    app.config["TESTING"] = True
    app.config["WTF_CSRF_ENABLED"] = False  # Disable CSRF for testing
    
    # Create all tables
    with app.app_context():
        Base.metadata.create_all(LocalSession().bind)
        yield app
        Base.metadata.drop_all(LocalSession().bind)


@pytest.fixture
def client(app):
    """Create a test client for the Flask application"""
    return app.test_client()


@pytest.fixture
def runner(app):
    """Create a test CLI runner for the Flask application"""
    return app.test_cli_runner()


# ------------------------------------------------
# Database Session Fixtures
# ------------------------------------------------

@pytest.fixture
def db_session(app):
    """Create a database session for testing"""
    session = LocalSession()
    yield session
    session.rollback()
    session.close()


# ------------------------------------------------
# User Fixtures
# ------------------------------------------------

@pytest.fixture
def test_user(db_session):
    """Create a test user"""
    user = User(
        username="testuser",
        email="test@example.com"
    )
    user.set_password("testpassword123")
    db_session.add(user)
    db_session.commit()
    db_session.refresh(user)
    return user


@pytest.fixture
def test_user2(db_session):
    """Create a second test user"""
    user = User(
        username="testuser2",
        email="test2@example.com"
    )
    user.set_password("testpassword123")
    db_session.add(user)
    db_session.commit()
    db_session.refresh(user)
    return user


@pytest.fixture
def authenticated_client(client, test_user):
    """Create an authenticated test client"""
    # Login by making a POST request to the login endpoint
    response = client.post("/auth/login", data={
        "email": test_user.email,
        "password": "testpassword123"
    }, follow_redirects=False)
    
    # Ensure login was successful
    assert response.status_code in [200, 302]
    
    return client


# ------------------------------------------------
# Document Fixtures
# ------------------------------------------------

@pytest.fixture
def test_document(db_session, test_user):
    """Create a test document"""
    document = Document(
        title="Test Document",
        content="This is test content for a document.",
        user_id=test_user.id
    )
    db_session.add(document)
    db_session.commit()
    db_session.refresh(document)
    return document


@pytest.fixture
def test_document_no_questions(db_session, test_user):
    """Create a test document without questions"""
    document = Document(
        title="Document Without Questions",
        content="This document has no questions yet.",
        user_id=test_user.id
    )
    db_session.add(document)
    db_session.commit()
    db_session.refresh(document)
    return document


# ------------------------------------------------
# Question Fixtures
# ------------------------------------------------

@pytest.fixture
def test_questions(db_session, test_document):
    """Create test questions for a document"""
    questions = []
    for i in range(3):
        question = Question(
            type=QuestionType.qcm,
            question=f"Test Question {i+1}?",
            choices=["Option A", "Option B", "Option C", "Option D"],
            answer="Option A",
            document_id=test_document.id
        )
        db_session.add(question)
        questions.append(question)
    db_session.commit()
    for q in questions:
        db_session.refresh(q)
    return questions


# ------------------------------------------------
# Mock Fixtures for External Dependencies
# ------------------------------------------------

@pytest.fixture
def mock_generate_mcq():
    """Mock the generate_mcq function"""
    mock_questions = [
        {
            "question": "What is the capital of France?",
            "answers": ["Paris", "London", "Berlin", "Madrid"],
            "correct_answer": "Paris"
        },
        {
            "question": "What is 2+2?",
            "answers": ["3", "4", "5", "6"],
            "correct_answer": "4"
        }
    ]
    
    with patch("app.routes.quizzes.generate_mcq", return_value=mock_questions):
        yield mock_questions


@pytest.fixture
def mock_generate_mcq_empty():
    """Mock the generate_mcq function to return empty list"""
    with patch("app.routes.quizzes.generate_mcq", return_value=[]):
        yield


@pytest.fixture
def mock_extract_text():
    """Mock the extract_text_from_file function"""
    with patch("app.routes.documents.extract_text_from_file", return_value="Extracted text content"):
        yield "Extracted text content"


# ------------------------------------------------
# File Upload Fixtures
# ------------------------------------------------

@pytest.fixture
def sample_pdf_file(tmp_path):
    """Create a sample PDF file for testing"""
    from reportlab.pdfgen import canvas
    file_path = tmp_path / "sample.pdf"
    c = canvas.Canvas(str(file_path))
    c.drawString(100, 750, "Sample PDF Content for Testing")
    c.save()
    return file_path


@pytest.fixture
def sample_docx_file(tmp_path):
    """Create a sample DOCX file for testing"""
    from docx import Document
    file_path = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("Sample DOCX Content for Testing")
    doc.save(file_path)
    return file_path

